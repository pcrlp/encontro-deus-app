"""
Gestão Encontro com Deus — Streamlit + Supabase
Versão Final - Etiquetas Perfeitas + Botão Atualizar na Secretaria
"""
import streamlit as st
import pandas as pd
import uuid, io, re, math, requests, unicodedata
from datetime import datetime, date, timezone
from supabase import create_client, Client
import streamlit.components.v1 as components

# ─── Config & Supabase ───────────────────────────────────────────────────────
def get_sb() -> Client:
    return create_client(st.secrets["SUPABASE_URL"], st.secrets["SUPABASE_SERVICE_KEY"])

def utcnow() -> str:
    return datetime.now(timezone.utc).isoformat()

# ─── Custom CSS (Light Mode / Clean SaaS) ────────────────────────────────────
def inject_custom_css():
    st.markdown("""
    <style>
    /* Força textos para cor escura e fundo principal claro */
    .stApp { background-color: #f8fafc; color: #1e293b; }
    
    /* FORÇA TEXTO ESCURO (Resolve o problema do celular em Modo Escuro) */
    p, span, h1, h2, h3, h4, h5, h6, label, li {
        color: #1e293b !important;
    }

    /* Neumorphic/SaaS Containers - Light */
    div[data-testid="stContainer"] {
        border: 1px solid #e2e8f0 !important;
        border-radius: 10px !important;
        background-color: #ffffff !important;
        box-shadow: 0 1px 3px 0 rgba(0, 0, 0, 0.1);
        padding: 1rem;
        margin-bottom: 1rem;
    }
    
    /* Ajuste para as Métricas do Dashboard */
    [data-testid="stMetricValue"] > div { color: #0f172a !important; }
    [data-testid="stMetricLabel"] * { color: #475569 !important; }

    /* Primary Buttons - Blue */
    .stButton>button[kind="primary"] {
        background-color: #2563eb !important;
        border: none !important;
        border-radius: 6px !important;
    }
    .stButton>button[kind="primary"] * {
        color: #ffffff !important;
        font-weight: 600 !important;
    }
    
    /* Download Buttons - Green */
    .stDownloadButton>button {
        background-color: #16a34a !important;
        border: none !important;
        border-radius: 6px !important;
    }
    .stDownloadButton>button * {
        color: #ffffff !important;
        font-weight: 600 !important;
    }
    
    /* Expanders */
    .streamlit-expanderHeader {
        background-color: #f1f5f9 !important;
        border-radius: 6px;
    }
    .streamlit-expanderHeader * {
        color: #0f172a !important;
    }
    
    /* Avisos (Success, Info, Warning) para não ficarem apagados */
    div[data-testid="stAlert"] * { color: #0f172a !important; }
    </style>
    """, unsafe_allow_html=True)

def open_multiple_links(links):
    if not links: return
    js_code = f"""
    <script>
    const links = {links};
    links.forEach(link => {{ window.open(link, '_blank'); }});
    </script>
    """
    components.html(js_code, height=0)

# ─── Enums ────────────────────────────────────────────────
GENDER_MAP = {0: "-", 1: "Masculino", 2: "Feminino"}
GENDER_REV = {"Não informado": 0, "Masculino": 1, "Feminino": 2}
SHIRT_MAP = {0: "-", 1: "PP", 2: "P", 3: "M", 4: "G", 5: "GG", 6: "G1", 7: "G2", 8: "G3", 9: "G4"}
SHIRT_REV = {v: k for k, v in SHIRT_MAP.items()}
SHIRT_KEYS = ["P", "M", "G", "GG", "G1", "G2", "G3", "G4"]
CATEGORY_OPTIONS = ["Encontrista", "Servo", "Servo (sem aquisição de camisa)", "Equipe"]

# ─── Helpers ──────────────────────────────────────────────────────────────────
def fmt_date_br(d):
    if not d: return ""
    try:
        parts = str(d).split("T")[0].split("-")
        return f"{parts[2]}/{parts[1]}/{parts[0]}"
    except: return str(d)

def norm(s):
    if not s: return ""
    return unicodedata.normalize("NFD", str(s)).encode("ascii", "ignore").decode().strip().lower()

def is_encounterist(cat): return "encontr" in norm(cat)
def is_server(cat): return "servo" in norm(cat)
def is_server_no_shirt(cat):
    c = norm(cat); return is_server(cat) and ("sem aquisicao" in c or "sem camisa" in c)
def is_server_with_shirt(cat): return is_server(cat) and not is_server_no_shirt(cat)

def age_from(birth_str):
    if not birth_str: return None
    try:
        d = str(birth_str).split("T")[0]; parts = d.split("-")
        y, m, day = int(parts[0]), int(parts[1]), int(parts[2])
        today = date.today(); age = today.year - y
        if (today.month, today.day) < (m, day): age -= 1
        return abs(age)
    except: return None

def parse_gender(val):
    if pd.isna(val): return 0
    s = str(val).strip().lower()
    if s in ("m", "masculino", "male", "1"): return 1
    if s in ("f", "feminino", "female", "2"): return 2
    return 0

def parse_shirt(val):
    if pd.isna(val): return 0
    return SHIRT_REV.get(str(val).strip().upper(), 0)

def safe_str(val, max_len=None):
    if pd.isna(val): return None
    s = str(val).strip()
    if not s or s == "nan": return None
    if max_len: s = s[:max_len]
    return s

def parse_date_br(val):
    if pd.isna(val): return None
    s = str(val).strip()
    for fmt in ("%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y", "%m/%d/%Y", "%d/%m/%y"):
        try: return datetime.strptime(s, fmt).date().isoformat()
        except: continue
    return None

# ─── Data loaders ─────────────────────────────────────────────────────────────
def load_events(): return get_sb().table("Events").select("*").order("CreatedAtUtc", desc=True).execute().data or []
def load_event(eid):
    r = get_sb().table("Events").select("*").eq("Id", eid).execute(); return r.data[0] if r.data else None
def load_participants(eid): return get_sb().table("Participants").select("*").eq("EventId", eid).order("Name").execute().data or []
def load_rooms(eid): return get_sb().table("Rooms").select("*").eq("EventId", eid).order("Name").execute().data or []
def load_assignments(eid): return get_sb().table("RoomAssignments").select("*").eq("EventId", eid).execute().data or []

def save_secretary_state(eid, team, dist, sec_status):
    import json
    payload = json.dumps({"team": team, "dist": dist, "status": sec_status})
    try: get_sb().table("Events").update({"SecretaryState": payload, "UpdatedAtUtc": utcnow()}).eq("Id", eid).execute()
    except: pass

def load_secretary_state(eid):
    import json
    try:
        ev = load_event(eid)
        if ev and ev.get("SecretaryState"):
            data = json.loads(ev["SecretaryState"])
            return data.get("team", []), data.get("dist", {}), data.get("status", {})
    except: pass
    return [], {}, {}

# ─── CSV Import ───────────────────────────────────────────────────────────────
def find_header(columns, candidates):
    for c in candidates:
        for col in columns:
            if norm(col) == norm(c): return col
    return None

def import_csv(event_id, file_bytes, replace=False):
    sb = get_sb()
    if replace:
        sb.table("RoomAssignments").delete().eq("EventId", event_id).execute()
        sb.table("Participants").delete().eq("EventId", event_id).execute()
    try: text = file_bytes.decode("utf-8-sig")
    except:
        try: text = file_bytes.decode("latin-1")
        except: text = file_bytes.decode("cp1252", errors="replace")
    first_line = text.split("\n")[0]
    delim = ";" if first_line.count(";") > first_line.count(",") else ","
    df = pd.read_csv(io.StringIO(text), sep=delim, dtype=str); df.columns = [c.strip() for c in df.columns]

    col_name    = find_header(df.columns, ["Nome do Encontrista", "Encontrista", "Nome", "Participante"])
    col_cat     = find_header(df.columns, ["Categoria"])
    col_email   = find_header(df.columns, ["Email", "E-mail"])
    col_gender  = find_header(df.columns, ["Sexo"])
    col_shirt   = find_header(df.columns, ["Tamanho da Camisa", "Tamanho da Camisa ", "Camiseta", "Camisa"])
    col_birth   = find_header(df.columns, ["Data de Nascimento", "Nascimento"])
    col_phone   = find_header(df.columns, ["Celular", "Telefone", "Celular (WhatsApp)", "Celular/WhatsApp", "Telefone Celular", "Celular / WhatsApp", "Contato (celular)"])
    col_marital = find_header(df.columns, ["Estado civil", "Estado Civil"])
    col_sector  = find_header(df.columns, ["Qual nome do setor do grupo de conexão?", "Qual o nome do setor do grupo de conexão?", "Setor de Conexão", "Setor de Conexao", "Setor (Centro, Norte...)", "Setor"])
    col_group   = find_header(df.columns, ["Nome do Grupo de Conexão", "Se a sua resposta anterior foi afirmativa, qual é o nome do seu GC?", "Qual é o nome do seu GC", "nome do seu gc", "Grupo de Conexão", "Grupo de Conexao", "Nome do GC", "GC"])
    col_invited = find_header(df.columns, ["Você veio a convite de alguém? Se sim, poderia me informar o nome da pessoa que o(a) convidou?", "Voce veio a convite de alguem", "Quem convidou", "Convidado por", "Indicado por", "Indicador por", "Quem indicou"])

    if not col_name: raise ValueError(f"Coluna 'Nome' não encontrada. Colunas: {', '.join(df.columns)}")
    ok, fail = 0, 0
    for _, row in df.iterrows():
        name = safe_str(row.get(col_name) if col_name else None, 200)
        if not name: fail += 1; continue
        record = {"Id": str(uuid.uuid4()), "EventId": event_id, "Name": name,
            "Email": safe_str(row.get(col_email) if col_email else None, 200),
            "Gender": parse_gender(row.get(col_gender) if col_gender else None),
            "BirthDate": parse_date_br(row.get(col_birth) if col_birth else None),
            "Phone": safe_str(row.get(col_phone) if col_phone else None, 40),
            "ShirtSize": parse_shirt(row.get(col_shirt) if col_shirt else None),
            "Category": safe_str(row.get(col_cat) if col_cat else None, 120),
            "ConnectionSector": safe_str(row.get(col_sector) if col_sector else None, 120),
            "ConnectionGroup": safe_str(row.get(col_group) if col_group else None, 120),
            "InvitedBy": safe_str(row.get(col_invited) if col_invited else None, 200),
            "CreatedAtUtc": utcnow()}
        record = {k: v for k, v in record.items() if v is not None}
        try: sb.table("Participants").insert(record).execute(); ok += 1
        except: fail += 1
    return ok, fail

# ─── Auto-distribute rooms ───────────────────────────────────────────────────
def distribute_rooms(event_id):
    sb = get_sb()
    sb.table("RoomAssignments").delete().eq("EventId", event_id).execute()
    all_rooms = load_rooms(event_id); parts = load_participants(event_id)
    if not all_rooms: return 0, 0, "Nenhum quarto cadastrado."
    enc = [p for p in parts if is_encounterist(p.get("Category"))]
    if not enc: return 0, 0, "Nenhum encontrista encontrado."

    def distrib(gv):
        rooms   = sorted([r for r in all_rooms if r.get("Gender") == gv], key=lambda r: r["Capacity"])
        people  = [p for p in enc if p.get("Gender") == gv]
        if not rooms or not people: return []
        seniors = sorted([p for p in people if (age_from(p.get("BirthDate")) or 0) >= 50], key=lambda p: -(age_from(p.get("BirthDate")) or 0))
        others  = sorted([p for p in people if (age_from(p.get("BirthDate")) or 0) <  50], key=lambda p: -(age_from(p.get("BirthDate")) or 0))
        total_cap = sum(r["Capacity"] for r in rooms); total_p = len(people)
        if total_cap <= 0: return []
        target = {r["Id"]: int(math.floor(total_p * r["Capacity"] / total_cap)) for r in rooms}
        rem = total_p - sum(target.values())
        order = sorted(rooms, key=lambda r: -(r["Capacity"] - target[r["Id"]])); idx = 0
        while rem > 0: target[order[idx % len(order)]["Id"]] += 1; rem -= 1; idx += 1
        assigns = []
        for room in rooms:
            need = min(room["Capacity"], target[room["Id"]])
            while need > 0 and seniors:
                assigns.append({"Id": str(uuid.uuid4()), "EventId": event_id, "RoomId": room["Id"], "ParticipantId": seniors.pop(0)["Id"], "CreatedAtUtc": utcnow()}); need -= 1
            while need > 0 and others:
                assigns.append({"Id": str(uuid.uuid4()), "EventId": event_id, "RoomId": room["Id"], "ParticipantId": others.pop(0)["Id"], "CreatedAtUtc": utcnow()}); need -= 1
        return assigns

    inserts = distrib(2) + distrib(1)
    for a in inserts: sb.table("RoomAssignments").insert(a).execute()
    return len(inserts), len(all_rooms), None

def generate_sector_pdf(participants, ev_name="Encontro com Deus", assigns=None, rooms=None):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    assigns = assigns or []; rooms = rooms or []
    rm = {r["Id"]: r["Name"] for r in rooms}
    pr = {a["ParticipantId"]: rm.get(a["RoomId"], "-") for a in assigns}

    SC = {"azul": "#DBEAFE", "amarelo": "#FEF9C3", "verde": "#DCFCE7", "lilas": "#F3E8FF", "vermelho": "#FEE2E2"}
    def sk(s):
        n = norm(s)
        for k in SC:
            if n.startswith(k): return k
        return ""
    ORDER = ["azul", "amarelo", "verde", "lilas", "vermelho", ""]

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=30, rightMargin=30, topMargin=36, bottomMargin=36)
    styles    = getSampleStyleSheet()
    cell_style = ParagraphStyle("cell", fontName="Helvetica",      fontSize=8, leading=10)
    hdr_style  = ParagraphStyle("hdr",  fontName="Helvetica-Bold", fontSize=8, leading=10)
    title_style = ParagraphStyle("title2", fontName="Helvetica-Bold", fontSize=13, leading=16, spaceAfter=4)
    sub_style   = ParagraphStyle("sub",   fontName="Helvetica", fontSize=9, leading=11, textColor=colors.HexColor("#555555"))

    def build_section(group_label, group_parts):
        elems = []
        elems.append(Paragraph(ev_name, title_style))
        elems.append(Paragraph(f"{group_label} — {len(group_parts)} pessoa(s)", sub_style))
        elems.append(Spacer(1, 10))

        sorted_parts = sorted(group_parts, key=lambda p: (ORDER.index(sk(p.get("ConnectionSector", ""))), p.get("Name", "")))
        data = [[ Paragraph("Nº", hdr_style), Paragraph("Nome", hdr_style), Paragraph("Categoria", hdr_style),
            Paragraph("Quarto", hdr_style), Paragraph("Setor / GC", hdr_style), Paragraph("Convidado por",hdr_style),
        ]]
        row_colors = []

        for i, p in enumerate(sorted_parts):
            sector = p.get("ConnectionSector") or ""; gc = p.get("ConnectionGroup") or "-"
            setor_gc = f"{sector or '-'} / {gc}" if sector else gc
            quarto = pr.get(p["Id"], "-"); invited = p.get("InvitedBy") or "-"; cat = p.get("Category") or "-"
            data.append([
                Paragraph(str(i+1), cell_style), Paragraph(p.get("Name", ""), cell_style), Paragraph(cat, cell_style),
                Paragraph(quarto, cell_style), Paragraph(setor_gc, cell_style), Paragraph(invited, cell_style),
            ])
            row_colors.append(colors.HexColor(SC.get(sk(sector), "#F9FAFB")))

        t = Table(data, colWidths=[22, 145, 75, 65, 118, 110], repeatRows=1)
        table_style = [
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#E2E8F0")), ("GRID", (0,0), (-1,-1), 0.4, colors.HexColor("#CBD5E1")),
            ("VALIGN", (0,0), (-1,-1), "TOP"), ("PADDING", (0,0), (-1,-1), 3),
        ]
        for i, c in enumerate(row_colors): table_style.append(("BACKGROUND", (0, i+1), (-1, i+1), c))
        t.setStyle(TableStyle(table_style))
        elems.append(t)
        return elems

    encontristas = [p for p in participants if is_encounterist(p.get("Category"))]
    servos       = [p for p in participants if is_server(p.get("Category"))]
    equipe       = [p for p in participants if norm(p.get("Category","")).startswith("equipe")]
    outros       = [p for p in participants if not is_encounterist(p.get("Category")) and not is_server(p.get("Category")) and not norm(p.get("Category","")).startswith("equipe")]

    all_elems = []; sections = []
    if encontristas: sections.append(("Encontristas", encontristas))
    if servos:       sections.append(("Servos", servos))
    if equipe:       sections.append(("Equipe", equipe))
    if outros:       sections.append(("Outros", outros))
    if not sections: sections = [("Participantes", participants)]

    for idx, (label, grp) in enumerate(sections):
        if idx > 0: all_elems.append(PageBreak())
        all_elems.extend(build_section(label, grp))

    doc.build(all_elems)
    return buf.getvalue()

def generate_rooms_pdf(event_id):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    rooms   = load_rooms(event_id)
    assigns = load_assignments(event_id)
    parts   = load_participants(event_id)
    pm      = {p["Id"]: p for p in parts}
    abr     = {}
    for a in assigns: abr.setdefault(a["RoomId"], []).append(a)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()

    cell_style = ParagraphStyle("rcell", fontName="Helvetica", fontSize=8, leading=10)
    hdr_style  = ParagraphStyle("rhdr", fontName="Helvetica-Bold", fontSize=8, leading=10)
    COL_W = [210, 32, 40, 145, 96]

    elems = []
    for ri, room in enumerate(rooms):
        if ri > 0: elems.append(PageBreak())
        leader = pm.get(room.get("LeaderId")) if room.get("LeaderId") else None

        elems.append(Paragraph(f"Quarto: {room['Name']}", styles["Title"]))
        elems.append(Paragraph(
            f"Sexo: {GENDER_MAP.get(room.get('Gender', 0), '-')} · Capacidade: {room['Capacity']} · Líder: {leader['Name'] if leader else '-'}",
            styles["Normal"]))
        elems.append(Spacer(1, 12))

        data = [[
            Paragraph("Nome", hdr_style), Paragraph("Idade", hdr_style), Paragraph("Camisa", hdr_style),
            Paragraph("GC", hdr_style), Paragraph("Setor", hdr_style),
        ]]

        for a in abr.get(room["Id"], []):
            p = pm.get(a["ParticipantId"])
            if not p: continue
            age = age_from(p.get("BirthDate"))
            data.append([
                Paragraph(p["Name"], cell_style), Paragraph(str(age) if age else "-", cell_style),
                Paragraph(SHIRT_MAP.get(p.get("ShirtSize", 0), "-"), cell_style),
                Paragraph(p.get("ConnectionGroup") or "-", cell_style), Paragraph(p.get("ConnectionSector") or "-", cell_style),
            ])

        t = Table(data, colWidths=COL_W, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F1F5F9")), ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1,-1), 8), ("GRID", (0, 0), (-1,-1), 0.5, colors.HexColor("#CBD5E1")),
            ("VALIGN", (0, 0), (-1,-1), "TOP"), ("PADDING", (0, 0), (-1,-1), 3),
        ]))
        elems.append(t)

    doc.build(elems)
    return buf.getvalue()

# ─── MÓDULO ETIQUETAS PIMACO (CENTRALIZAÇÃO E EMPILHAMENTO PERFEITOS) ─────────
def generate_labels_pimaco(parts_sel, label_type, assigns, rooms):
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas as rl_canvas
    import math as _math

    MM = 2.8346
    rm = {r["Id"]: r["Name"] for r in rooms}
    pr = {a["ParticipantId"]: rm.get(a["RoomId"], "-") for a in assigns}

    PAGE_W, PAGE_H = letter # Papel Carta
    LBL_W = 66.7 * MM
    LBL_H = 25.4 * MM
    COLS = 3
    ROWS = 10
    PER_PAGE = COLS * ROWS

    MARGIN_LEFT = 7.9 * MM
    MARGIN_TOP = PAGE_H - (12.7 * MM)

    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=letter)
    total = len(parts_sel)
    pages = _math.ceil(total / PER_PAGE) if total else 1

    def split_text_to_lines(text, font, size, max_width):
        c.setFont(font, size)
        if c.stringWidth(text, font, size) <= max_width: return [text]
        words = text.split()
        if len(words) <= 1: return [text]
        
        mid = len(words) // 2 + (len(words) % 2)
        l1 = " ".join(words[:mid])
        l2 = " ".join(words[mid:])
        return [l1, l2]

    for pg in range(pages):
        batch = parts_sel[pg * PER_PAGE:(pg + 1) * PER_PAGE]
        for idx, p in enumerate(batch):
            col = idx % COLS
            row = idx // COLS
            
            x = MARGIN_LEFT + col * LBL_W
            y = MARGIN_TOP - (row + 1) * LBL_H

            cx = x + (LBL_W / 2)
            cy = y + (LBL_H / 2)
            
            # Margem super segura para os lados (desconta 3mm de cada lado)
            max_w = LBL_W - (6 * MM) 

            # Deixa sempre o nome em MAIÚSCULO para padronizar o visual
            name = str(p.get("Name", "")).strip().upper()

            if label_type == "nome":
                name_size = 10
                name_font = "Helvetica-Bold"
                lines = split_text_to_lines(name, name_font, name_size, max_w)
                
                while name_size > 6 and any(c.stringWidth(l, name_font, name_size) > max_w for l in lines):
                    name_size -= 0.5
                    lines = split_text_to_lines(name, name_font, name_size, max_w)

                total_h = len(lines) * (name_size * 1.2)
                start_y = cy + (total_h / 2) - (name_size * 0.8)
                
                c.setFont(name_font, name_size)
                for l in lines:
                    c.drawCentredString(cx, start_y, l)
                    start_y -= (name_size * 1.2)

            else:
                quarto = pr.get(p["Id"], "SEM QUARTO")
                
                if label_type == "blusa":
                    shirt = SHIRT_MAP.get(p.get("ShirtSize", 0), "-")
                    info_lines = [f"Quarto: {quarto}", f"Camisa: {shirt}"]
                else:
                    cat = str(p.get("Category") or "-").upper()
                    info_lines = [cat[:35], f"Quarto: {quarto}"]

                name_size = 9
                name_font = "Helvetica-Bold"
                lines = split_text_to_lines(name, name_font, name_size, max_w)
                
                while name_size > 6 and any(c.stringWidth(l, name_font, name_size) > max_w for l in lines):
                    name_size -= 0.5
                    lines = split_text_to_lines(name, name_font, name_size, max_w)

                info_size = 8
                info_font = "Helvetica"
                
                name_line_h = name_size * 1.2
                info_line_h = info_size * 1.2
                
                # Calculando o bloco completo para encaixar exatamente no meio da etiqueta
                total_h = (len(lines) * name_line_h) + (len(info_lines) * info_line_h)
                
                # A partir do centro (cy), sobe metade do tamanho do bloco inteiro
                start_y = cy + (total_h / 2) - (name_size * 0.8)
                
                c.setFont(name_font, name_size)
                for l in lines:
                    c.drawCentredString(cx, start_y, l)
                    start_y -= name_line_h
                    
                c.setFont(info_font, info_size)
                for info in info_lines:
                    c.drawCentredString(cx, start_y, info)
                    start_y -= info_line_h

        if pg < pages - 1: c.showPage()

    c.save()
    return buf.getvalue()


# ─── Google Sheets Integration ───────────────────────────────────────────────
def sheets_url_to_csv(url):
    if "export?format=csv" in url.lower(): return url
    m = re.search(r"spreadsheets/d/([A-Za-z0-9\-_]+)", url, re.I)
    if m:
        sid = m.group(1); gid = "0"
        gm  = re.search(r"[?#&]gid=([0-9]+)", url, re.I)
        if gm: gid = gm.group(1)
        return f"https://docs.google.com/spreadsheets/d/{sid}/export?format=csv&id={sid}&gid={gid}"
    return url

def fetch_sheet_csv(url):
    csv_url = sheets_url_to_csv(url)
    try:
        resp = requests.get(csv_url, timeout=30, headers={"User-Agent": "Mozilla/5.0"})
        if resp.status_code == 200 and len(resp.text) > 10: return resp.text
    except: pass
    m = re.search(r"spreadsheets/d/([A-Za-z0-9\-_]+)", url, re.I)
    if m:
        sid = m.group(1); gid = "0"
        gm = re.search(r"[?#&]gid=([0-9]+)", url, re.I)
        if gm: gid = gm.group(1)
        gviz = f"https://docs.google.com/spreadsheets/d/{sid}/gviz/tq?tqx=out:csv&gid={gid}"
        resp2 = requests.get(gviz, timeout=30, headers={"User-Agent": "Mozilla/5.0"})
        if resp2.status_code == 200: return resp2.text
    raise Exception("Não foi possível carregar a planilha. Verifique se o link está público.")

def _do_load_letters(eid, url):
    if not url: return False, "URL vazia"
    try:
        text = fetch_sheet_csv(url)
        df = pd.read_csv(io.StringIO(text), dtype=str); df.columns = [c.strip() for c in df.columns]
        
        ct = find_header(df.columns, ["Para quem","Destinatário","Para","Nome do Encontrista","Encontrista","Nome do encontrista:","Nome do encontrista"])
        cf = find_header(df.columns, ["De quem","Remetente","De","Nome de quem escreve","Seu nome","Seu nome:"])
        cm = find_header(df.columns, ["Mensagem","Carta","Texto","Conteúdo","Escreva algo especial para ele(a) aqui:","Escreva algo especial para ele(a) aqui","Mensagem para ele(a)"])
        
        if not ct or not cm: 
            return False, f"Colunas necessárias não encontradas nas cartas. Colunas atuais: {', '.join(df.columns)}"
            
        letters = {}
        for _, row in df.iterrows():
            to = safe_str(row.get(ct))
            sender = safe_str(row.get(cf)) if cf else "—"
            if sender is None: sender = "—"
            msg = safe_str(row.get(cm)) or ""
            if to and msg: letters.setdefault(to, []).append({"sender": sender, "message": msg})
            
        total = sum(len(v) for v in letters.values())
        st.session_state[f"letters_dl_{eid}"] = {"ts": datetime.now().strftime("%H:%M:%S"), "total": total}
        st.session_state[f"letters_data_{eid}"] = letters
        return True, f"✅ {total} carta(s) carregadas"
    except Exception as e: return False, str(e)

def _do_load_photos(eid, url):
    if not url: return False, "URL vazia"
    try:
        text = fetch_sheet_csv(url)
        df = pd.read_csv(io.StringIO(text), dtype=str); df.columns = [c.strip() for c in df.columns]
        
        cn = find_header(df.columns, ["Nome do Encontrista", "Nome do encontrista:", "Nome do encontrista", "Encontrista", "Nome"])
        cp = find_header(df.columns, ["Foto", "Fotos", "URL da Foto", "URL das Fotos", "URL", "Link"])
        
        if not cn or not cp: 
            return False, f"Colunas necessárias não encontradas nas fotos. Colunas atuais: {', '.join(df.columns)}"
            
        groups = {}
        for _, row in df.iterrows():
            name = safe_str(row.get(cn)); photo = safe_str(row.get(cp))
            if name and photo:
                groups.setdefault(name, [])
                for lk in photo.split(","):
                    lk = lk.strip()
                    if lk: groups[name].append(lk)
                    
        total = sum(len(v) for v in groups.values())
        st.session_state[f"photos_dl_{eid}"] = {"ts": datetime.now().strftime("%H:%M:%S"), "total": total}
        st.session_state[f"photo_groups_{eid}"] = groups
        return True, f"✅ {total} foto(s) mapeadas"
    except Exception as e: return False, str(e)

def make_gdrive_view_url(raw_url):
    if not raw_url: return None
    raw_url = raw_url.strip(); fid = None
    m = re.search(r"[?&]id=([A-Za-z0-9\-_]+)", raw_url, re.I)
    if m: fid = m.group(1)
    if not fid:
        m = re.search(r"/d/([A-Za-z0-9\-_]+)", raw_url, re.I)
        if m: fid = m.group(1)
    if not fid and re.match(r"^[A-Za-z0-9\-_]{20,}$", raw_url): fid = raw_url
    if fid: return f"https://drive.google.com/file/d/{fid}/view"
    return raw_url

def generate_letters_docx(participant_name, letters_list):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)
    
    for i, carta in enumerate(letters_list):
        if i > 0: 
            doc.add_page_break()
        
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_h = h.add_run(f"Para: {participant_name}")
        run_h.bold = True
        run_h.font.size = Pt(14)
        doc.add_paragraph()
        
        sender_p = doc.add_paragraph()
        run_s = sender_p.add_run(f"De: {carta.get('sender','—')}")
        run_s.bold = True
        run_s.font.size = Pt(11)
        doc.add_paragraph()
        
        msg_p = doc.add_paragraph(carta.get("message",""))
        if msg_p.runs: 
            msg_p.runs[0].font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

# ═══════════════════════════════════════════════════════════════════════════════
# PAGES
# ═══════════════════════════════════════════════════════════════════════════════
def show_login():
    st.markdown("## 🔐 Gestão Encontro com Deus")
    st.caption("Informe a senha para acessar.")
    pwd = st.text_input("Senha", type="password", key="lp")
    if st.button("Entrar", type="primary", use_container_width=True):
        if pwd == st.secrets.get("APP_PASSWORD", "encontro2025"):
            st.session_state.authenticated = True
            st.query_params["auth"] = "1"
            st.rerun()
        else: st.error("Senha incorreta.")

def page_events():
    st.markdown("## ⛪ Eventos"); events = load_events()
    _, cr = st.columns([4, 1])
    with cr:
        if st.button("＋ Criar novo evento", type="primary", use_container_width=True): st.session_state.page = "event_new"; st.rerun()
    if not events: st.info("Nenhum evento criado ainda."); return
    for ev in events:
        with st.container():
            a, b, c = st.columns([4, 2, 1])
            with a: st.markdown(f"**{ev['Name']}**"); st.caption(f"{fmt_date_br(ev.get('StartDate'))} • {fmt_date_br(ev.get('EndDate'))}")
            with b:
                if st.button("Abrir Dashboard →", key=f"d_{ev['Id']}", use_container_width=True, type="primary"): 
                    st.session_state.current_event = ev["Id"]; st.session_state.page = "dashboard"; st.rerun()
            with c:
                if st.button("🗑️", key=f"x_{ev['Id']}"):
                    sb = get_sb()
                    for t in ["RoomAssignments","ScheduleItems","Rooms","Participants"]: sb.table(t).delete().eq("EventId", ev["Id"]).execute()
                    sb.table("Events").delete().eq("Id", ev["Id"]).execute(); st.rerun()

def page_event_new():
    st.markdown("## Novo Evento")
    if st.button("← Voltar"): st.session_state.page = "events"; st.rerun()
    with st.form("ne"):
        name = st.text_input("Nome do Evento", placeholder="Encontro com Deus - Abril")
        c1, c2 = st.columns(2)
        with c1: start = st.date_input("Data Início")
        with c2: end   = st.date_input("Data Fim")
        if st.form_submit_button("Salvar e abrir Dashboard", type="primary"):
            if not name or len(name) < 3: st.error("Nome mín. 3 caracteres."); return
            nid = str(uuid.uuid4())
            get_sb().table("Events").insert({"Id": nid, "Name": name, "StartDate": start.isoformat(), "EndDate": end.isoformat(), "Status": "Config", "CreatedAtUtc": utcnow()}).execute()
            st.session_state.current_event = nid; st.session_state.page = "dashboard"; st.rerun()

def event_sidebar(eid):
    ev = load_event(eid)
    if not ev: st.error("Evento não encontrado."); st.session_state.page = "events"; st.rerun(); return None
    with st.sidebar:
        st.markdown(f"### {ev['Name']}")
        if ev.get("StartDate"): st.caption(f"📅 {fmt_date_br(ev['StartDate'])} — {fmt_date_br(ev.get('EndDate'))}")
        st.divider()
        menu_items = {
            "dashboard":"📊 Dashboard",
            "participants":"👥 Participantes",
            "rooms":"🏠 Quartos",
            "labels":"🏷️ Etiquetas",
            "letters":"💌 Cartas",
            "photos":"📸 Fotos",
            "secretary":"🗂️ Secretaria",
            "print_management":"🖨️ Gestão de Impressão",
            "checkin_status":"✅ Acomp. Check-in",
            "settings":"⚙️ Configurações"
        }
        for k, label in menu_items.items():
            is_active = st.session_state.get("page") == k
            if st.button(label, key=f"n_{k}", use_container_width=True, type="primary" if is_active else "secondary"): 
                st.session_state.page = k; st.rerun()
        st.divider()
        if st.button("← Todos os eventos", use_container_width=True): st.session_state.pop("current_event", None); st.session_state.page = "events"; st.rerun()
    return ev

def page_dashboard(eid, ev):
    st.markdown(f"## 📊 Dashboard — {ev['Name']}")
    parts   = load_participants(eid)
    enc     = [p for p in parts if is_encounterist(p.get("Category"))]
    srv_com = [p for p in parts if is_server_with_shirt(p.get("Category"))]
    srv_sem = [p for p in parts if is_server_no_shirt(p.get("Category"))]
    equipe  = [p for p in parts if norm(p.get("Category","")).startswith("equipe")]
    
    c1,c2,c3,c4 = st.columns(4)
    c1.metric("Encontristas (M)", sum(1 for p in enc if p.get("Gender")==1))
    c2.metric("Encontristas (F)", sum(1 for p in enc if p.get("Gender")==2))
    c3.metric("Servos c/ camisa", len(srv_com)); c4.metric("Servos s/ camisa", len(srv_sem))
    c5,c6,c7,c8 = st.columns(4)
    c5.metric("Equipe", len(equipe)); c6.metric("Total Geral", len(parts)); c7.metric("Quartos", len(load_rooms(eid))); c8.metric("","")

    st.divider()
    col_title, col_btn = st.columns([3, 1])
    with col_title:
        st.markdown("### 📈 Progresso de Recebimentos")
    with col_btn:
        if st.button("🔄 Atualizar Dados Agora", type="primary", use_container_width=True):
            has_error = False
            with st.spinner("Buscando no Google Drive..."):
                if ev.get("LettersSheetUrl"): 
                    ok, msg = _do_load_letters(eid, ev["LettersSheetUrl"])
                    if not ok: st.error(f"Erro nas Cartas: {msg}"); has_error = True
                if ev.get("PhotosSheetUrl"): 
                    ok, msg = _do_load_photos(eid, ev["PhotosSheetUrl"])
                    if not ok: st.error(f"Erro nas Fotos: {msg}"); has_error = True
            if not has_error: st.rerun()

    letters = st.session_state.get(f"letters_data_{eid}", {})
    photo_groups = st.session_state.get(f"photo_groups_{eid}", {})

    def has_data(name, data_dict):
        return any(norm(name) == norm(k) or norm(name) in norm(k) or norm(k) in norm(name) for k in data_dict.keys())

    total_enc = len(enc)
    if total_enc > 0:
        faltam_cartas = [p for p in enc if not has_data(p["Name"], letters)]
        faltam_fotos = [p for p in enc if not has_data(p["Name"], photo_groups)]

        recebeu_carta = total_enc - len(faltam_cartas)
        nao_recebeu_carta = len(faltam_cartas)
        
        recebeu_foto = total_enc - len(faltam_fotos)
        nao_recebeu_foto = len(faltam_fotos)

        _, col_chart1, col_chart2, _ = st.columns([1, 4, 4, 1])
        
        try:
            import plotly.express as px
            with col_chart1:
                fig_cartas = px.pie(names=["Com Cartas", "Faltam"], values=[recebeu_carta, nao_recebeu_carta],
                                    title=f"Cartas (Base: {total_enc})", color_discrete_sequence=["#10b981", "#e2e8f0"], hole=0.6)
                fig_cartas.update_traces(textposition='inside', textinfo='percent+value', hoverinfo='label+percent')
                fig_cartas.update_layout(height=280, showlegend=True, legend=dict(orientation="h", yanchor="bottom", y=-0.3, xanchor="center", x=0.5), margin=dict(t=30, b=0, l=0, r=0))
                st.plotly_chart(fig_cartas, use_container_width=True)
            with col_chart2:
                fig_fotos = px.pie(names=["Com Fotos", "Faltam"], values=[recebeu_foto, nao_recebeu_foto],
                                   title=f"Fotos (Base: {total_enc})", color_discrete_sequence=["#3b82f6", "#e2e8f0"], hole=0.6)
                fig_fotos.update_traces(textposition='inside', textinfo='percent+value', hoverinfo='label+percent')
                fig_fotos.update_layout(height=280, showlegend=True, legend=dict(orientation="h", yanchor="bottom", y=-0.3, xanchor="center", x=0.5), margin=dict(t=30, b=0, l=0, r=0))
                st.plotly_chart(fig_fotos, use_container_width=True)
        except ImportError:
            st.warning("Gráficos de pizza indisponíveis. Instale a biblioteca 'plotly'.")

        st.divider()
        col_tab1, col_tab2 = st.columns(2)
        with col_tab1:
            st.markdown("🚨 **Encontristas SEM Cartas**")
            if faltam_cartas: st.dataframe(pd.DataFrame([{"Nome": p["Name"], "GC": p.get("ConnectionGroup") or "-"} for p in faltam_cartas]), hide_index=True, use_container_width=True)
            else: st.success("Todos os encontristas já receberam cartas!")
        with col_tab2:
            st.markdown("🚨 **Encontristas SEM Fotos**")
            if faltam_fotos: st.dataframe(pd.DataFrame([{"Nome": p["Name"], "GC": p.get("ConnectionGroup") or "-"} for p in faltam_fotos]), hide_index=True, use_container_width=True)
            else: st.success("Todos os encontristas já receberam fotos!")
    else:
        st.info("Nenhum encontrista cadastrado para gerar os gráficos.")

def page_participants(eid, ev):
    st.markdown(f"## 👥 Participantes — {ev['Name']}")
    with st.expander("📥 Importar CSV / ➕ Adicionar", expanded=False):
        uploaded = st.file_uploader("Arquivo CSV", type=["csv"], key="cu")
        replace  = st.checkbox("Substituir existentes", value=False)
        if st.button("Importar", type="primary") and uploaded:
            with st.spinner("Importando..."):
                ok, fail = import_csv(eid, uploaded.read(), replace); st.success(f"✅ {ok} importados, {fail} falha(s)."); st.rerun()
        st.divider()
        with st.form("ap"):
            a1,a2 = st.columns(2)
            with a1: pn=st.text_input("Nome*"); pg=st.selectbox("Gênero",["Masculino","Feminino","Não informado"]); pp=st.text_input("Telefone")
            with a2: pc=st.selectbox("Categoria",CATEGORY_OPTIONS); ps=st.selectbox("Camiseta",["-"]+SHIRT_KEYS); psc=st.text_input("Setor"); pgc=st.text_input("Grupo de Conexão")
            if st.form_submit_button("Adicionar Manualmente", type="primary"):
                if not pn: st.error("Nome obrigatório."); return
                get_sb().table("Participants").insert({"Id":str(uuid.uuid4()),"EventId":eid,"Name":pn,"Gender":GENDER_REV.get(pg,0),"ShirtSize":SHIRT_REV.get(ps,0),"Category":pc,"Phone":pp or None,"ConnectionSector":psc or None,"ConnectionGroup":pgc or None,"CreatedAtUtc":utcnow()}).execute()
                st.success(f"'{pn}' adicionado!"); st.rerun()

    parts  = load_participants(eid); assigns = load_assignments(eid); rooms = load_rooms(eid)
    rm = {r["Id"]: r["Name"] for r in rooms}; pr = {a["ParticipantId"]: rm.get(a["RoomId"],"-") for a in assigns}

    fc1,fc2,fc3,fc4 = st.columns(4)
    with fc1: search = st.text_input("🔍 Buscar Nome", key="ps")
    with fc2: fg     = st.selectbox("Sexo", ["Todos","Masculino","Feminino"], key="fg")
    with fc3: fc_sel = st.selectbox("Categoria", ["Todas","Encontrista","Servo","Equipe"], key="fc")
    with fc4: fsc    = st.text_input("Setor", key="fsc")
    
    filtered = parts
    if search:       filtered = [p for p in filtered if search.lower() in p["Name"].lower()]
    if fg=="Masculino":  filtered = [p for p in filtered if p.get("Gender")==1]
    elif fg=="Feminino": filtered = [p for p in filtered if p.get("Gender")==2]
    if fc_sel=="Encontrista": filtered = [p for p in filtered if is_encounterist(p.get("Category"))]
    elif fc_sel=="Servo":     filtered = [p for p in filtered if is_server(p.get("Category"))]
    elif fc_sel=="Equipe":    filtered = [p for p in filtered if norm(p.get("Category","")).startswith("equipe")]
    if fsc: filtered = [p for p in filtered if fsc.lower() in norm(p.get("ConnectionSector",""))]

    st.markdown(f"**{len(filtered)} participante(s)**")
    ac1, ac2 = st.columns(2)
    with ac1:
        if filtered:
            pdf_data = generate_sector_pdf(filtered, ev_name=ev["Name"], assigns=assigns, rooms=rooms)
            st.download_button("⬇️ Baixar PDF Setores", data=pdf_data, file_name="setores.pdf", mime="application/pdf", use_container_width=True)
    with ac2:
        if filtered:
            rows_exp = [{"Nome": p["Name"], "Categoria": p.get("Category") or "-", "Sexo": GENDER_MAP.get(p.get("Gender",0),"-"), "Camisa": SHIRT_MAP.get(p.get("ShirtSize",0),"-"), "Quarto": pr.get(p["Id"],"-"), "Setor": p.get("ConnectionSector") or "-"} for p in filtered]
            csv_data = pd.DataFrame(rows_exp).to_csv(index=False).encode('utf-8')
            st.download_button("⬇️ Baixar CSV", data=csv_data, file_name="participantes.csv", mime="text/csv", use_container_width=True)

    st.divider()
    if not filtered: st.info("Nenhum participante encontrado.")
    else:
        for p in filtered:
            pid = p["Id"]; edit_key = f"pedit_{pid}"; quarto = pr.get(pid, "-")
            with st.container():
                c1, c2, c3 = st.columns([6, 1, 1])
                with c1:
                    st.markdown(f"**{p['Name']}** · {GENDER_MAP.get(p.get('Gender',0),'-')} · Camisa: {SHIRT_MAP.get(p.get('ShirtSize',0),'-')}")
                    st.caption(f"{p.get('Category') or '-'} · Quarto: {quarto} · Setor: {p.get('ConnectionSector') or '-'}")
                with c2:
                    if st.button("✏️ Editar", key=f"ebtn_{pid}", use_container_width=True): st.session_state[edit_key] = not st.session_state.get(edit_key, False); st.rerun()
                with c3:
                    if st.button("🗑️ Apagar", key=f"delbtn_{pid}", use_container_width=True):
                        sb2 = get_sb(); sb2.table("RoomAssignments").delete().eq("ParticipantId", pid).execute(); sb2.table("Participants").delete().eq("Id", pid).execute(); st.rerun()
                
                if st.session_state.get(edit_key, False):
                    with st.form(f"pform_{pid}"):
                        ea, eb = st.columns(2)
                        with ea:
                            en    = st.text_input("Nome", value=p.get("Name",""))
                            eg    = st.selectbox("Gênero", ["Masculino","Feminino","Não informado"], index=max(0, p.get("Gender",0)-1) if p.get("Gender",0)>0 else 2)
                            ecat  = st.selectbox("Categoria", CATEGORY_OPTIONS, index=CATEGORY_OPTIONS.index(p["Category"]) if p.get("Category") in CATEGORY_OPTIONS else 0)
                        with eb:
                            esh  = st.selectbox("Camisa", ["-"]+SHIRT_KEYS, index=(["-"]+SHIRT_KEYS).index(SHIRT_MAP.get(p.get("ShirtSize",0),"-")))
                            esc  = st.text_input("Setor", value=p.get("ConnectionSector") or "")
                            egc  = st.text_input("GC", value=p.get("ConnectionGroup") or "")
                        if st.form_submit_button("💾 Salvar Modificações", type="primary"):
                            get_sb().table("Participants").update({"Name": en, "Gender": GENDER_REV.get(eg, 0), "ShirtSize": SHIRT_REV.get(esh, 0), "Category": ecat, "ConnectionSector": esc or None, "ConnectionGroup": egc or None, "UpdatedAtUtc": utcnow()}).eq("Id", pid).execute()
                            st.session_state.pop(edit_key, None); st.rerun()

def page_rooms(eid, ev):
    st.markdown(f"## 🏠 Quartos — {ev['Name']}")
    parts  = load_participants(eid); rooms = load_rooms(eid); assigns = load_assignments(eid)
    pm     = {p["Id"]: p for p in parts}; sb2 = get_sb()

    b1, b2, b3 = st.columns([1, 1, 2])
    with b1:
        if st.button("🔀 Auto-Distribuir", type="primary", use_container_width=True):
            with st.spinner("Distribuindo..."):
                n, nr, err = distribute_rooms(eid)
                if err: st.error(err)
                else: st.success(f"✅ {n} alocados!"); st.rerun()
    with b2:
        if rooms:
            pdf_r = generate_rooms_pdf(eid)
            st.download_button("⬇️ Baixar PDF Quartos", data=pdf_r, file_name="quartos.pdf", mime="application/pdf", use_container_width=True)

    with st.expander("➕ Novo Quarto", expanded=False):
        with st.form("nr"):
            rc1,rc2,rc3 = st.columns(3)
            with rc1: rn   = st.text_input("Nome")
            with rc2: rcap = st.number_input("Capacidade", min_value=1, value=10)
            with rc3: rg   = st.selectbox("Gênero", ["Feminino","Masculino"])
            if st.form_submit_button("Criar Quarto", type="primary"):
                if rn: sb2.table("Rooms").insert({"Id":str(uuid.uuid4()),"EventId":eid,"Name":rn,"Capacity":rcap,"Gender":2 if rg=="Feminino" else 1,"CreatedAtUtc":utcnow()}).execute(); st.rerun()

    st.divider()
    if not rooms: st.info("Nenhum quarto cadastrado."); return

    am = {}; aids = set()
    for a in assigns: 
        am.setdefault(a["RoomId"], []).append(a)
        aids.add(a["ParticipantId"])

    for room in rooms:
        rid = room["Id"]; occs = am.get(rid, [])
        gender_str = GENDER_MAP.get(room.get('Gender', 0), '-')
        
        with st.expander(f"🏠 {room['Name']} ({gender_str}) — {len(occs)}/{room['Capacity']} vagas", expanded=False):
            col_ed1, col_ed2 = st.columns([3, 1])
            with col_ed1:
                with st.popover("⚙️ Editar Dados do Quarto"):
                    new_rn = st.text_input("Nome", value=room['Name'], key=f"rn_{rid}")
                    new_cap = st.number_input("Capacidade", value=room['Capacity'], min_value=1, key=f"rc_{rid}")
                    new_g_str = "Masculino" if room.get('Gender')==1 else "Feminino"
                    new_g = st.selectbox("Gênero", ["Feminino","Masculino"], index=["Feminino","Masculino"].index(new_g_str), key=f"rg_{rid}")
                    if st.button("💾 Salvar Alterações", type="primary", key=f"sroom_{rid}"):
                        sb2.table("Rooms").update({"Name":new_rn, "Capacity":new_cap, "Gender":2 if new_g=="Feminino" else 1}).eq("Id",rid).execute(); st.rerun()
            with col_ed2:
                if st.button("🗑️ Apagar Quarto", key=f"delr_{rid}", use_container_width=True):
                    sb2.table("RoomAssignments").delete().eq("RoomId",rid).execute(); sb2.table("Rooms").delete().eq("Id",rid).execute(); st.rerun()

            st.divider()

            if occs:
                for a in occs:
                    p = pm.get(a["ParticipantId"])
                    if not p: continue
                    o1, o2, o3 = st.columns([5, 3, 1])
                    with o1: st.markdown(f"👤 **{p['Name']}**")
                    with o2:
                        other_rooms = [r for r in rooms if r["Id"] != rid and (r.get("Gender",0)==p.get("Gender",0) or r.get("Gender",0)==0)]
                        if other_rooms:
                            with st.popover("🔄 Trocar/Mover"):
                                dest_room_name = st.selectbox("Mover para:", [r["Name"] for r in other_rooms], key=f"dest_{a['Id']}")
                                dest_room = next(r for r in other_rooms if r["Name"]==dest_room_name)
                                dest_occs = am.get(dest_room["Id"], [])
                                
                                if len(dest_occs) >= dest_room["Capacity"]:
                                    st.warning("O quarto destino está cheio. Com quem você quer trocar?")
                                    swap_candidates = [pm.get(da["ParticipantId"]) for da in dest_occs if pm.get(da["ParticipantId"])]
                                    swap_p = st.selectbox("Trocar lugar com:", [c["Name"] for c in swap_candidates], key=f"swapc_{a['Id']}")
                                    if st.button("Executar Troca", type="primary", key=f"bswap_{a['Id']}"):
                                        try:
                                            asgn_a_id = a["Id"]
                                            target_pid = next((c["Id"] for c in swap_candidates if c["Name"] == swap_p), None)
                                            target_asgn = next((da for da in dest_occs if da["ParticipantId"] == target_pid), None)
                                            
                                            if target_asgn:
                                                asgn_b_id = target_asgn["Id"]
                                                sb2.table("RoomAssignments").update({"RoomId": rid}).eq("Id", asgn_b_id).execute()
                                                sb2.table("RoomAssignments").update({"RoomId": dest_room["Id"]}).eq("Id", asgn_a_id).execute()
                                                st.rerun()
                                        except Exception as e: st.error(f"Erro ao realizar troca: {e}")
                                else:
                                    if st.button("Mover para Quarto", type="primary", key=f"bmove_{a['Id']}"):
                                        try:
                                            sb2.table("RoomAssignments").update({"RoomId": dest_room["Id"]}).eq("Id", a["Id"]).execute()
                                            st.rerun()
                                        except Exception as e: st.error(f"Erro ao mover: {e}")
                    with o3:
                        if st.button("✕ Remover", key=f"rem_{a['Id']}"):
                            sb2.table("RoomAssignments").delete().eq("Id",a["Id"]).execute(); st.rerun()
            
            unass = [p for p in parts if p["Id"] not in aids and is_encounterist(p.get("Category")) and (p.get("Gender",0)==room.get("Gender",0) or room.get("Gender",0)==0)]
            if unass and len(occs) < room["Capacity"]:
                sp = st.selectbox("Adicionar pessoa sem quarto:", [""]+[p["Name"] for p in unass], key=f"addp_{rid}")
                if sp and st.button("＋ Adicionar", type="primary", key=f"badd_{rid}"):
                    pid = next(p["Id"] for p in unass if p["Name"]==sp)
                    sb2.table("RoomAssignments").insert({"Id":str(uuid.uuid4()),"EventId":eid,"RoomId":rid,"ParticipantId":pid,"CreatedAtUtc":utcnow()}).execute(); st.rerun()

# ─── MÓDULO ETIQUETAS ───────────────────────────────
def page_labels(eid, ev):
    st.markdown(f"## 🏷️ Etiquetas — {ev['Name']}")
    st.caption("Modelo: **Pimaco 6180** · Carta · 66,7 × 25,4 mm · **30 etiquetas/folha**")
    st.warning("⚠️ **MUITO IMPORTANTE:** Na hora de imprimir o PDF, verifique nas configurações da impressora se a 'Escala' está em **100%** ou **Tamanho Real**. Não use 'Ajustar à Página'.")

    parts = load_participants(eid); assigns = load_assignments(eid); rooms = load_rooms(eid)
    rm = {r["Id"]: r["Name"] for r in rooms}; pr = {a["ParticipantId"]: rm.get(a["RoomId"],"-") for a in assigns}

    f1, f2, f3 = st.columns(3)
    with f1: search_lbl = st.text_input("🔍 Buscar por nome", key="lbl_search")
    with f2: cat_lbl = st.selectbox("Categoria", ["Todos","Encontrista","Servo","Equipe"], key="lbl_cat")
    with f3: room_opts = ["Todos os quartos"] + sorted([r["Name"] for r in rooms]); room_lbl = st.selectbox("🏠 Quarto", room_opts, key="lbl_room")

    filtered_lbl = parts
    if search_lbl: filtered_lbl = [p for p in filtered_lbl if search_lbl.lower() in p["Name"].lower()]
    if cat_lbl == "Encontrista": filtered_lbl = [p for p in filtered_lbl if is_encounterist(p.get("Category"))]
    elif cat_lbl == "Servo": filtered_lbl = [p for p in filtered_lbl if is_server(p.get("Category"))]
    elif cat_lbl == "Equipe": filtered_lbl = [p for p in filtered_lbl if norm(p.get("Category","")).startswith("equipe")]
    if room_lbl != "Todos os quartos": filtered_lbl = [p for p in filtered_lbl if pr.get(p["Id"], "-") == room_lbl]

    repeat_qty = st.number_input("🔁 Cópias por pessoa", min_value=1, max_value=20, value=1, step=1, help="Ex: 3 = João João João Maria...")
    
    flag_key = f"lbl_flags_{eid}"
    if flag_key not in st.session_state: st.session_state[flag_key] = set()
    
    st.divider()

    col_sel1, col_sel2, col_sel3 = st.columns([2, 2, 6])
    with col_sel1:
        if st.button("☑️ Marcar Todos", use_container_width=True):
            st.session_state[flag_key] = set(p["Id"] for p in filtered_lbl)
            st.rerun()
    with col_sel2:
        if st.button("☐ Desmarcar Todos", use_container_width=True):
            st.session_state[flag_key] = set()
            st.rerun()

    parts_flagged = [p for p in filtered_lbl if p["Id"] in st.session_state[flag_key]]
    parts_to_print = [p for p in parts_flagged for _ in range(int(repeat_qty))]
    
    if parts_to_print:
        st.info(f"📄 **{len(parts_to_print)}** etiqueta(s) selecionada(s) para impressão.")
        b1, b2 = st.columns(2)
        with b1:
            pdf_c = generate_labels_pimaco(parts_to_print, "blusa", assigns, rooms)
            st.download_button("🖨️ Baixar Etiquetas (Com Blusa/Quarto)", pdf_c, "etiquetas_completas.pdf", "application/pdf", use_container_width=True, type="primary")
        with b2:
            pdf_n = generate_labels_pimaco(parts_to_print, "nome", assigns, rooms)
            st.download_button("🖨️ Baixar Etiquetas (Só Nome)", pdf_n, "etiquetas_nomes.pdf", "application/pdf", use_container_width=True)
    else:
        st.warning("Nenhuma etiqueta selecionada. Marque os participantes abaixo.")

    st.markdown(f"**Lista de Participantes ({len(filtered_lbl)})**")
    for p in filtered_lbl:
        pid = p["Id"]
        is_checked = pid in st.session_state[flag_key]
        
        c_flag, c_name = st.columns([1, 9])
        with c_flag:
            ui_check = st.checkbox("", value=is_checked, key=f"ui_{pid}_{is_checked}", label_visibility="collapsed")
            if ui_check != is_checked:
                if ui_check: st.session_state[flag_key].add(pid)
                else: st.session_state[flag_key].discard(pid)
                st.rerun()
        with c_name:
            st.markdown(f"**{p['Name']}** - {pr.get(pid, '-')} - Camisa: {SHIRT_MAP.get(p.get('ShirtSize',0),'-')}")

# ─── MÓDULO ABA DE CARTAS ───────────────────────────────
def page_letters(eid, ev):
    st.markdown(f"## 💌 Cartas — {ev['Name']}")
    parts = load_participants(eid)
    enc = [p for p in parts if is_encounterist(p.get("Category"))]

    col_s, col_b = st.columns([3, 1])
    with col_s:
        search_l = st.text_input("🔍 Buscar por nome", key="search_letters")
    with col_b:
        if st.button("🔄 Atualizar Cartas", type="primary", use_container_width=True):
            if ev.get("LettersSheetUrl"):
                ok, msg = _do_load_letters(eid, ev["LettersSheetUrl"])
                if ok: st.success("Cartas atualizadas!")
                else: st.error(f"Erro: {msg}")
            else: st.warning("URL não configurada na aba Configurações.")

    if search_l:
        enc = [p for p in enc if search_l.lower() in p["Name"].lower()]

    letters_dict = st.session_state.get(f"letters_data_{eid}", {})
    
    if not letters_dict and not ev.get("LettersSheetUrl"):
        st.info("As cartas não estão configuradas. Adicione o link em 'Configurações'.")
        return

    st.divider()
    for p in enc:
        user_letters = []
        for key, lts in letters_dict.items():
            if norm(key) == norm(p["Name"]) or p["Name"].lower() in key.lower() or key.lower() in p["Name"].lower():
                user_letters.extend(lts)
        
        with st.container():
            c1, c2 = st.columns([5, 2])
            c1.markdown(f"**{p['Name']}**")
            c1.caption(f"Total: {len(user_letters)} carta(s) recebida(s)")
            
            if user_letters:
                docx_bytes = generate_letters_docx(p["Name"], user_letters)
                c2.download_button(
                    label=f"⬇️ Baixar {len(user_letters)} Cartas (Word)", 
                    data=docx_bytes, 
                    file_name=f"Cartas_{norm(p['Name'])}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"ltr_dl_{p['Id']}",
                    use_container_width=True
                )
            else:
                c2.button("Sem cartas", disabled=True, key=f"ltr_dsb_{p['Id']}", use_container_width=True)

# ─── MÓDULO ABA DE FOTOS ───────────────────────────────
def page_photos(eid, ev):
    st.markdown(f"## 📸 Fotos — {ev['Name']}")
    st.info("⚠️ Para abrir todas as fotos de uma vez, permita os **Pop-ups** no seu navegador.")
    parts = load_participants(eid)
    enc = [p for p in parts if is_encounterist(p.get("Category"))]

    col_s, col_b = st.columns([3, 1])
    with col_s:
        search_f = st.text_input("🔍 Buscar por nome", key="search_photos")
    with col_b:
        if st.button("🔄 Atualizar Fotos", type="primary", use_container_width=True):
            if ev.get("PhotosSheetUrl"):
                ok, msg = _do_load_photos(eid, ev["PhotosSheetUrl"])
                if ok: st.success("Fotos atualizadas!")
                else: st.error(f"Erro: {msg}")
            else: st.warning("URL não configurada na aba Configurações.")

    if search_f:
        enc = [p for p in enc if search_f.lower() in p["Name"].lower()]

    photo_groups = st.session_state.get(f"photo_groups_{eid}", {})

    if not photo_groups and not ev.get("PhotosSheetUrl"):
        st.info("As fotos não estão configuradas. Adicione o link em 'Configurações'.")
        return

    st.divider()
    for p in enc:
        user_photos = []
        for key, lks in photo_groups.items():
            if norm(key) == norm(p["Name"]) or p["Name"].lower() in key.lower() or key.lower() in p["Name"].lower():
                user_photos.extend(lks)
        
        with st.container():
            c1, c2 = st.columns([5, 2])
            c1.markdown(f"**{p['Name']}**")
            c1.caption(f"Total: {len(user_photos)} link(s) de foto(s)")
            
            if user_photos:
                links_urls = [make_gdrive_view_url(lk) or lk for lk in user_photos]
                if c2.button(f"🚀 Abrir {len(user_photos)} Fotos (Nova Aba)", key=f"pht_op_{p['Id']}", use_container_width=True, type="primary"):
                    open_multiple_links(links_urls)
            else:
                c2.button("Sem fotos", disabled=True, key=f"pht_dsb_{p['Id']}", use_container_width=True)

# ─── MÓDULO SECRETARIA ───────────────────────────────
def page_secretary(eid, ev):
    st.markdown(f"## 🗂️ Secretaria — {ev['Name']}")
    
    col_title, col_btn = st.columns([3, 1])
    with col_title:
        st.write("") # Espaço vazio para alinhar o botão à direita
    with col_btn:
        if st.button("🔄 Atualizar Cartas/Fotos", type="primary", use_container_width=True, key="btn_update_sec"):
            has_error = False
            with st.spinner("Buscando no Google Drive..."):
                if ev.get("LettersSheetUrl"):
                    ok, msg = _do_load_letters(eid, ev["LettersSheetUrl"])
                    if not ok: st.error(f"Erro nas Cartas: {msg}"); has_error = True
                if ev.get("PhotosSheetUrl"):
                    ok, msg = _do_load_photos(eid, ev["PhotosSheetUrl"])
                    if not ok: st.error(f"Erro nas Fotos: {msg}"); has_error = True
            if not has_error: st.rerun()

    parts  = load_participants(eid); assigns = load_assignments(eid); rooms = load_rooms(eid)
    enc    = [p for p in parts if is_encounterist(p.get("Category"))]
    rm     = {r["Id"]: r["Name"] for r in rooms}; pr = {a["ParticipantId"]: rm.get(a["RoomId"], "-") for a in assigns}
    
    db_team, db_dist, db_status = load_secretary_state(eid)
    team = st.session_state.get(f"sec_team_{eid}", db_team)
    dist = st.session_state.get(f"sec_dist_{eid}", db_dist)
    sec_status = st.session_state.get(f"sec_status_{eid}", db_status)

    def persist(): 
        save_secretary_state(eid, team, dist, sec_status)
        st.session_state[f"sec_team_{eid}"] = team
        st.session_state[f"sec_dist_{eid}"] = dist
        st.session_state[f"sec_status_{eid}"] = sec_status

    letters = st.session_state.get(f"letters_data_{eid}", {})
    photo_groups = st.session_state.get(f"photo_groups_{eid}", {})

    def count_letters(name):
        cnt = 0
        for key, lts in letters.items():
            if norm(key)==norm(name) or name.lower() in key.lower() or key.lower() in name.lower(): cnt += len(lts)
        return cnt

    def count_photos(name):
        for key, links in photo_groups.items():
            if norm(key)==norm(name) or name.lower() in key.lower() or key.lower() in name.lower(): return len(links)
        return 0

    tab1, tab2 = st.tabs(["📋 Acompanhamento das Bolsas", "⚙️ Distribuição de Equipe"])
    
    with tab1:
        if not dist:
            st.warning("Vá na aba 'Distribuição de Equipe' para atribuir as bolsas aos responsáveis.")
        else:
            membro_sel = st.selectbox("👤 Ver acompanhamento de:", list(dist.keys()), key="sec_member_sel")
            pids_sel = dist.get(membro_sel, []); people_sel = [p for p in enc if p["Id"] in pids_sel]
            
            # --- Barra de Progresso Acompanhamento ---
            total_bags = len(people_sel)
            closed_bags = sum(1 for p in people_sel if sec_status.get(p["Id"], {}).get("bolsa_ok", False))
            
            st.markdown(f"**Progresso de {membro_sel}:** {closed_bags} de {total_bags} bolsas finalizadas.")
            st.progress(closed_bags / total_bags if total_bags > 0 else 0)
            st.divider()
            
            for p in people_sel:
                pid = p["Id"]; ps = sec_status.get(pid, {})
                bolsa_ok = ps.get("bolsa_ok", False); cartas_ok = ps.get("cartas_ok", False); fotos_ok = ps.get("fotos_ok", False)
                print_stat = ps.get("print_status", "none")
                n_cartas = count_letters(p["Name"]); n_fotos = count_photos(p["Name"])
                quarto = pr.get(pid, "-")
                
                with st.container():
                    r1, r2, r3 = st.columns([3, 2, 2])
                    with r1:
                        st.markdown(f"**{p['Name']}**")
                        st.caption(f"Quarto: {quarto}")
                    with r2:
                        st.markdown(f"💌 **Cartas:** {n_cartas} | 📷 **Fotos:** {n_fotos}")
                        if not bolsa_ok:
                            c1, c2 = st.columns(2)
                            with c1:
                                if st.checkbox("Cartas OK", value=cartas_ok, key=f"ck_{pid}") != cartas_ok:
                                    if pid not in sec_status: sec_status[pid] = {}
                                    sec_status[pid]["cartas_ok"] = not cartas_ok; persist(); st.rerun()
                            with c2:
                                if st.checkbox("Fotos OK", value=fotos_ok, key=f"fk_{pid}") != fotos_ok:
                                    if pid not in sec_status: sec_status[pid] = {}
                                    sec_status[pid]["fotos_ok"] = not fotos_ok; persist(); st.rerun()
                    with r3:
                        if bolsa_ok:
                            st.success("✅ Bolsa Finalizada")
                            # Botões de Reabrir ou Reiniciar Completamente
                            s1, s2 = st.columns(2)
                            with s1:
                                if st.button("↩️ Reabrir", key=f"reopen_{pid}", help="Apenas reabre a bolsa mantendo os marcadores"):
                                    sec_status[pid]["bolsa_ok"] = False; persist(); st.rerun()
                            with s2:
                                if st.button("🔄 Reiniciar", key=f"reset1_{pid}", help="Zera o status completamente"):
                                    sec_status[pid] = {"bolsa_ok": False, "cartas_ok": False, "fotos_ok": False, "print_status": "none"}
                                    persist(); st.rerun()
                        else:
                            if print_stat == "requested":
                                st.warning("🖨️ Aguardando Impressão")
                                if st.button("🔄 Reiniciar / Cancelar Solicitação", key=f"reset2_{pid}"):
                                    sec_status[pid] = {"bolsa_ok": False, "cartas_ok": False, "fotos_ok": False, "print_status": "none"}
                                    persist(); st.rerun()
                            elif print_stat == "printing":
                                st.info("🖨️ Sendo impresso (Letícia)")
                            elif print_stat == "done":
                                if st.button("✅ Fechar Bolsa", type="primary", key=f"done_{pid}", use_container_width=True):
                                    sec_status[pid]["bolsa_ok"] = True; persist(); st.rerun()
                            else:
                                if st.button("🖨️ Solicitar Impressão", key=f"reqprint_{pid}", use_container_width=True):
                                    if pid not in sec_status: sec_status[pid] = {}
                                    sec_status[pid]["print_status"] = "requested"
                                    sec_status[pid]["print_req_by"] = membro_sel
                                    persist(); st.rerun()

    with tab2:
        st.markdown("#### Equipe Responsável pelas Bolsas")
        servos = [p for p in parts if is_server(p.get("Category"))]; servo_names = [p["Name"] for p in servos]
        novo_membro = st.selectbox("Selecionar membro", [""] + servo_names, key="sec_sel_servo")
        if st.button("➕ Adicionar membro à equipe") and novo_membro:
            if novo_membro not in team: team.append(novo_membro); persist(); st.rerun()
            
        st.write(f"Membros cadastrados: {', '.join(team) if team else 'Nenhum'}")
        st.divider()
        if st.button("🔀 Distribuir encontristas automaticamente (Divisão igualitária)", type="primary"):
            by_room = {}
            for p in enc: quarto = pr.get(p["Id"], "Sem quarto"); by_room.setdefault(quarto, []).append(p)
            ordered = [item for sublist in [by_room[k] for k in sorted(by_room.keys())] for item in sublist]
            new_dist = {m: [] for m in team}
            for idx, p in enumerate(ordered): new_dist[team[idx % len(team)]].append(p["Id"])
            dist = new_dist; persist(); st.success("✅ Distribuição feita com sucesso!"); st.rerun()

# ─── MÓDULO DE ACOMPANHAMENTO DO CHECK-IN ─────────────────────────────────────
def page_checkin_status(eid, ev):
    st.markdown(f"## ✅ Acompanhamento de Check-in — {ev['Name']}")
    
    parts = load_participants(eid)
    enc = [p for p in parts if is_encounterist(p.get("Category", ""))]

    homens = [p for p in enc if p.get("Gender") == 1]
    mulheres = [p for p in enc if p.get("Gender") == 2]

    h_chegaram = sum(1 for h in homens if h.get("CheckInStatus"))
    h_faltam = len(homens) - h_chegaram

    m_chegaram = sum(1 for m in mulheres if m.get("CheckInStatus"))
    m_faltam = len(mulheres) - m_chegaram

    total_enc = len(enc)
    total_chegaram = h_chegaram + m_chegaram
    total_faltam = h_faltam + m_faltam

    st.markdown(f"""
    <style>
    .stats-container-chk {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 10px; margin: 0.5rem 0 1rem 0; }}
    .stats-box-chk {{ background: #FFFFFF; border: 1px solid #E2E8F0; border-radius: 12px; padding: 0.8rem; text-align: center; box-shadow: 0 1px 3px rgba(0,0,0,0.05); }}
    .stats-title-chk {{ font-size: 0.85rem; color: #475569; font-weight: 600; text-transform: uppercase; letter-spacing: 0.5px; margin-bottom: 0.5rem; border-bottom: 1px solid #F1F5F9; padding-bottom: 0.3rem; }}
    .stats-row-chk {{ display: flex; justify-content: space-between; margin-top: 0.3rem; }}
    .stat-item-chk {{ text-align: center; flex: 1; }}
    .stat-number-chk {{ font-size: 1.3rem; font-weight: 700; color: #0f172a; }}
    .stat-number-chk.arrived {{ color: #16A34A; }}
    .stat-number-chk.missing {{ color: #EF4444; }}
    .stat-label-chk {{ font-size: 0.65rem; color: #64748B; text-transform: uppercase; font-weight: 600; }}
    </style>
    <div class="stats-container-chk">
        <div class="stats-box-chk">
            <div class="stats-title-chk">👥 Total ({total_enc})</div>
            <div class="stats-row-chk">
                <div class="stat-item-chk"><div class="stat-number-chk arrived">{total_chegaram}</div><div class="stat-label-chk">Chegaram</div></div>
                <div class="stat-item-chk"><div class="stat-number-chk missing">{total_faltam}</div><div class="stat-label-chk">Faltam</div></div>
            </div>
        </div>
        <div class="stats-box-chk">
            <div class="stats-title-chk">👨 Homens ({len(homens)})</div>
            <div class="stats-row-chk">
                <div class="stat-item-chk"><div class="stat-number-chk arrived">{h_chegaram}</div><div class="stat-label-chk">Chegaram</div></div>
                <div class="stat-item-chk"><div class="stat-number-chk missing">{h_faltam}</div><div class="stat-label-chk">Faltam</div></div>
            </div>
        </div>
        <div class="stats-box-chk">
            <div class="stats-title-chk">👩 Mulheres ({len(mulheres)})</div>
            <div class="stats-row-chk">
                <div class="stat-item-chk"><div class="stat-number-chk arrived">{m_chegaram}</div><div class="stat-label-chk">Chegaram</div></div>
                <div class="stat-item-chk"><div class="stat-number-chk missing">{m_faltam}</div><div class="stat-label-chk">Faltam</div></div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.divider()
    st.markdown("### 🚨 Encontristas que ainda NÃO chegaram")

    search_faltantes = st.text_input("🔍 Buscar por nome...", key="search_faltantes")

    faltantes = [p for p in enc if not p.get("CheckInStatus")]

    if search_faltantes:
        faltantes = [p for p in faltantes if search_faltantes.lower() in p.get("Name", "").lower()]

    if faltantes:
        data = []
        for p in faltantes:
            data.append({
                "Nome": p.get("Name", ""),
                "Telefone": p.get("Phone") or "-",
                "Setor": p.get("ConnectionSector") or "-",
                "GC": p.get("ConnectionGroup") or "-",
                "Indicado por": p.get("InvitedBy") or "-"
            })
        st.dataframe(pd.DataFrame(data), hide_index=True, use_container_width=True)
    else:
        if search_faltantes:
            st.info("Nenhum encontrista faltante encontrado com esse nome.")
        else:
            st.success("🎉 Todos os encontristas já chegaram!")

# ─── MÓDULO CENTRAL DE IMPRESSÃO ───────────────────────────────
def page_print_management(eid, ev):
    st.markdown(f"## 🖨️ Gestão de Impressão — {ev['Name']}")
    st.info("⚠️ **Atenção (Equipe de Impressão):** Ao imprimir, vá imediatamente no Google Forms e **remova o nome do encontrista** para travar o recebimento de novas cartas ou fotos. Para abrir as fotos, libere os **Pop-ups** do navegador.")
    
    parts = load_participants(eid)
    _, _, sec_status = load_secretary_state(eid)
    
    print_queue = [p for p in parts if sec_status.get(p["Id"], {}).get("print_status") in ["requested", "printing"]]
    
    if not print_queue:
        st.success("Nenhuma solicitação pendente. A secretaria envia solicitações para esta tela.")
        return

    def mark_as_printing(pid):
        sec_status[pid]["print_status"] = "printing"
        save_secretary_state(eid, *load_secretary_state(eid)[:2], sec_status)

    for p in print_queue:
        pid = p["Id"]
        status = sec_status[pid]["print_status"]
        req_by = sec_status[pid].get("print_req_by", "Secretaria")
        
        with st.container():
            c1, c2, c3 = st.columns([4, 2, 2])
            with c1:
                st.markdown(f"### {p['Name']}")
                st.caption(f"Solicitante: **{req_by}** · Status: **{'🟡 Aguardando' if status == 'requested' else '🖨️ Em Andamento (Notificado)'}**")
            
            letters_dict = st.session_state.get(f"letters_data_{eid}", {})
            user_letters = []
            for key, lts in letters_dict.items():
                if norm(key) == norm(p["Name"]) or p["Name"].lower() in key.lower() or key.lower() in p["Name"].lower():
                    user_letters.extend(lts)
            
            with c2:
                if user_letters:
                    docx_bytes = generate_letters_docx(p["Name"], user_letters)
                    st.download_button(
                        label=f"⬇️ Baixar Cartas ({len(user_letters)})", 
                        data=docx_bytes, 
                        file_name=f"Cartas_{norm(p['Name'])}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        on_click=mark_as_printing, args=(pid,),
                        use_container_width=True, type="primary" if status=="requested" else "secondary"
                    )
                else:
                    st.button("Sem cartas", disabled=True, use_container_width=True)

            with c3:
                photo_groups = st.session_state.get(f"photo_groups_{eid}", {})
                user_photos = []
                for key, lks in photo_groups.items():
                    if norm(key) == norm(p["Name"]) or p["Name"].lower() in key.lower() or key.lower() in p["Name"].lower():
                        user_photos.extend(lks)
                
                if user_photos:
                    links_urls = [make_gdrive_view_url(lk) or lk for lk in user_photos]
                    if st.button(f"🚀 Abrir Fotos ({len(user_photos)})", key=f"open_prnt_{pid}", use_container_width=True, type="primary" if status=="requested" else "secondary"):
                        mark_as_printing(pid)
                        open_multiple_links(links_urls)
                        st.rerun()
                else:
                    st.button("Sem fotos", disabled=True, use_container_width=True)
            
            if status == "printing":
                st.divider()
                if st.button("✅ Concluir Impressão (Remover da Fila)", key=f"fin_{pid}", type="primary"):
                    sec_status[pid]["print_status"] = "done"
                    save_secretary_state(eid, *load_secretary_state(eid)[:2], sec_status)
                    st.rerun()

def page_settings(eid, ev):
    st.markdown(f"## ⚙️ Configurações — {ev['Name']}")
    with st.form("cfg"):
        name = st.text_input("Nome do Evento", value=ev.get("Name",""))
        lu = st.text_input("Link da Planilha de Cartas (Google Sheets)", value=ev.get("LettersSheetUrl") or "")
        phu = st.text_input("Link da Planilha de Fotos (Google Sheets)", value=ev.get("PhotosSheetUrl") or "")
        st.caption("Atenção: Garanta que o link seja de compartilhamento 'Qualquer pessoa com o link pode ler'")
        if st.form_submit_button("💾 Salvar Configurações", type="primary"):
            get_sb().table("Events").update({"Name":name, "LettersSheetUrl":lu or None, "PhotosSheetUrl":phu or None, "UpdatedAtUtc":utcnow()}).eq("Id",eid).execute()
            st.success("Configurações salvas!"); st.rerun()

# ═══════════════════════════════════════════════════════════════════════════════
def main():
    st.set_page_config(page_title="Gestão Encontro com Deus", page_icon="⛪", layout="wide", initial_sidebar_state="expanded")
    inject_custom_css()

    if "authenticated" not in st.session_state: st.session_state.authenticated = False
    if not st.session_state.authenticated and st.query_params.get("auth") == "1": st.session_state.authenticated = True

    if not st.session_state.authenticated: show_login(); return

    with st.sidebar:
        st.markdown("### ⛪ Encontro com Deus"); st.divider()
        if st.button("🚪 Sair", use_container_width=True):
            st.session_state.authenticated = False; st.query_params.clear(); st.session_state.pop("current_event", None); st.rerun()

    if "page" not in st.session_state: st.session_state.page = "events"
    page = st.session_state.page; eid = st.session_state.get("current_event")
    
    if page == "event_new": page_event_new(); return
    if page == "events" or not eid: page_events(); return
    
    ev = event_sidebar(eid)
    if not ev: return
    
    routes = {
        "dashboard": page_dashboard,
        "participants": page_participants,
        "rooms": page_rooms,
        "labels": page_labels,
        "letters": page_letters,
        "photos": page_photos,
        "secretary": page_secretary,
        "print_management": page_print_management,
        "checkin_status": page_checkin_status,
        "settings": page_settings
    }
    
    routes.get(page, page_dashboard)(eid, ev)

if __name__ == "__main__": main()
